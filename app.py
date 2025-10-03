import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

st.set_page_config(page_title="Reporte Juicios Evaluativos - SENA", layout="wide")

# --- APP PRINCIPAL ---
st.markdown("""
    <style>
    .main { background-color: #ffffff; }
    .stApp { background-color: #f0f8f4; }
    .info-card {
        background-color: #e8f5e9;
        padding: 15px;
        border-radius: 10px;
        border-left: 5px solid #2e7d32;
        margin: 10px 0;
    }
    .metric-card {
        background-color: #ffffff;
        padding: 15px;
        border-radius: 8px;
        border: 1px solid #e0e0e0;
        text-align: center;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    </style>
""", unsafe_allow_html=True)

# Header mejorado
col1, col2, col3 = st.columns([1, 2, 1])
with col2:
    st.markdown("""
    <div style='text-align: center;'>
        <h3 style='color: #2e7d32; margin-bottom: 5px;'>Centro Agroempresarial y Turístico de los Andes</h3>
        <h5 style='color: #555; margin-top: 0;'>Regional Santander</h5>
        <hr style='border: 2px solid #2e7d32; margin: 10px 0;'>
        <h4 style='color: #2e7d32;'>Reporte de Juicios Evaluativos - SENA</h4>
    </div>
    """, unsafe_allow_html=True)

st.markdown("---")

uploaded_file = st.file_uploader("📂 Sube el archivo Excel (.xls o .xlsx)", type=["xls", "xlsx"])


# --- MARCA DE AGUA (FOOTER) ---
st.markdown("---")
st.caption("Desarrollado por: Andrey y Ing. Faiver Adrian Abril")

if uploaded_file:
    try:
        df_raw = pd.read_excel(uploaded_file, header=None)
        ficha = str(df_raw.iloc[2, 2])
        denominacion = str(df_raw.iloc[5, 2])
        estado_ficha = str(df_raw.iloc[6, 2])
        fecha_inicio = pd.to_datetime(df_raw.iloc[7, 2]).date()
        fecha_fin = pd.to_datetime(df_raw.iloc[8, 2]).date()

        df = pd.read_excel(uploaded_file, skiprows=12)
        df = df[["Nombre", "Apellidos", "Competencia", "Resultado de Aprendizaje", "Juicio de Evaluación"]]

        # Información general compacta
        st.markdown("### 📊 Resumen Ficha")

        # Primera fila: Información básica
        col1, col2, col3, col4 = st.columns(4)

        with col1:
            st.metric("📋 Ficha", ficha)

        with col2:
            # Acortar denominación si es muy larga
            nombre_corto = denominacion[:20] + "..." if len(denominacion) > 20 else denominacion
            st.metric("🏷️ Programa", nombre_corto)

        with col3:
            # Estado con icono simple
            icono_estado = "🟢" if "ACTIV" in estado_ficha.upper() else "🟡" if "TERMIN" in estado_ficha.upper() else "🔴"
            st.metric(f"{icono_estado} Estado", estado_ficha)

        with col4:
            total_aprendices = len(df[['Nombre', 'Apellidos']].drop_duplicates())
            st.metric("👥 Aprendices", total_aprendices)

        # Segunda fila: Estadísticas y fechas
        col1, col2, col3 = st.columns(3)

        with col1:
            total_competencias = len(df['Competencia'].unique())
            st.metric("📚 Competencias", total_competencias)

        with col2:
            st.metric("📅 Inicio", fecha_inicio.strftime("%d/%m/%y"))

        with col3:
            st.metric("📅 Fin", fecha_fin.strftime("%d/%m/%y"))

        st.markdown("---")

        # Pestañas
        pestañas = st.tabs(["📊 Análisis General", "👥 Resumen por Ficha", "📤 Exportar Reportes", "❓ Ayuda"])

        # --- Análisis General ---
        with pestañas[0]:
            st.subheader("📊 Análisis General de Juicios Evaluativos")

            # Filtros en una sola línea
            col_filtro1, col_filtro2, col_espacio = st.columns([2, 2, 1])

            with col_filtro1:
                aprendices = df["Nombre"] + " " + df["Apellidos"]
                aprendiz_seleccionado = st.selectbox(
                    "👨‍🎓 Seleccionar Aprendiz",
                    ["Todos los aprendices"] + list(sorted(aprendices.unique())),
                    help="Filtrar por aprendiz específico"
                )

            with col_filtro2:
                # Aplicar filtro de aprendiz primero
                if aprendiz_seleccionado != "Todos los aprendices":
                    df_filtrado = df[aprendices == aprendiz_seleccionado]
                else:
                    df_filtrado = df

                competencias = df_filtrado["Competencia"].unique()
                competencia_seleccionada = st.selectbox(
                    "📚 Seleccionar Competencia",
                    ["Todas las competencias"] + list(sorted(competencias)),
                    help="Filtrar por competencia específica"
                )

            # Aplicar filtro de competencia
            if competencia_seleccionada != "Todas las competencias":
                df_filtrado = df_filtrado[df_filtrado["Competencia"] == competencia_seleccionada]

            # Mostrar resumen de filtros
            if not df_filtrado.empty:
                juicio_counts = df_filtrado["Juicio de Evaluación"].value_counts()

                # Layout principal: Gráfico + Estadísticas
                col_grafico, col_stats = st.columns([2, 1])

                with col_grafico:
                    # Gráfico de pastel mejorado
                    fig, ax = plt.subplots(figsize=(6, 4))
                    colors = ['#4CAF50', '#FF9800', '#F44336', '#2196F3', '#9C27B0']

                    wedges, texts, autotexts = ax.pie(
                        juicio_counts.values,
                        labels=juicio_counts.index,
                        autopct='%1.1f%%',
                        startangle=90,
                        colors=colors[:len(juicio_counts)],
                        textprops={'fontsize': 9, 'fontweight': 'bold'}
                    )

                    # Mejorar porcentajes
                    for autotext in autotexts:
                        autotext.set_color('white')
                        autotext.set_fontweight('bold')

                    ax.set_title('Distribución de Juicios', fontweight='bold', fontsize=12)
                    ax.axis('equal')
                    plt.tight_layout()
                    st.pyplot(fig)

                with col_stats:
                    st.markdown("### 📈 Estadísticas")
                    st.metric("Total Registros", len(df_filtrado))

                    # Mostrar distribución numérica
                    st.markdown("**Distribución:**")
                    for juicio, count in juicio_counts.items():
                        color = "#4CAF50" if "APROBADO" in juicio.upper() else "#FF9800" if "EVALUAR" in juicio.upper() else "#F44336"
                        st.markdown(f"<span style='color:{color}'>●</span> **{juicio}:** {count}",
                                unsafe_allow_html=True)


                # Vista previa de datos
                st.markdown("### 👁️ Registros Pendientes por Evaluar")

                # Filtrar solo los registros con juicio "POR EVALUAR"
                df_por_evaluar = df_filtrado[df_filtrado['Juicio de Evaluación'] == 'POR EVALUAR']

                if not df_por_evaluar.empty:
                    st.markdown(f"**Se encontraron {len(df_por_evaluar)} registros pendientes por evaluar:**")

                    # Mostrar tabla con los registros pendientes
                    st.dataframe(
                        df_por_evaluar,
                        use_container_width=True,
                        hide_index=True,
                        height=min(400, len(df_por_evaluar) * 35 + 40)  # Ajustar altura automáticamente
                    )

                    # Mostrar estadísticas adicionales
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Total Pendientes", len(df_por_evaluar))
                    with col2:
                        aprendices_pendientes = df_por_evaluar[['Nombre', 'Apellidos']].drop_duplicates().shape[0]
                        st.metric("Aprendices con Pendientes", aprendices_pendientes)
                    with col3:
                        competencias_pendientes = df_por_evaluar['Competencia'].nunique()
                        st.metric("Competencias Pendientes", competencias_pendientes)

                else:
                    st.success("🎉 **¡Excelente! No hay registros pendientes por evaluar con los filtros actuales.**")
                    st.info("""
                    **Todos los juicios evaluativos han sido completados para los criterios seleccionados:**
                    - Aprendiz: {aprendiz_seleccionado}
                    - Competencia: {competencia_seleccionada}
                    """.format(
                        aprendiz_seleccionado=aprendiz_seleccionado,
                        competencia_seleccionada=competencia_seleccionada
                    ))

            # Mostrar mensaje informativo si hay más registros en el filtro original
            if len(df_filtrado) > len(df_por_evaluar):
                otros_registros = len(df_filtrado) - len(df_por_evaluar)
                st.caption(f"💡 Además, hay {otros_registros} registros con juicios evaluativos completados (APROBADO/NO APROBADO).")


        # --- Resumen por Aprendiz ---
        with pestañas[1]:
            st.subheader("👥 Resumen por Ficha")

            # Calcular resumen
            resumen_data = []
            aprendices_unicos = df[['Nombre', 'Apellidos']].drop_duplicates()

            for _, row in aprendices_unicos.iterrows():
                nombre_completo = f"{row['Nombre']} {row['Apellidos']}"
                datos_aprendiz = df[(df['Nombre'] == row['Nombre']) & (df['Apellidos'] == row['Apellidos'])]

                por_evaluar = len(datos_aprendiz[datos_aprendiz['Juicio de Evaluación'].str.upper() == 'POR EVALUAR'])
                aprobado = len(datos_aprendiz[datos_aprendiz['Juicio de Evaluación'].str.upper() == 'APROBADO'])
                no_aprobado = len(datos_aprendiz[datos_aprendiz['Juicio de Evaluación'].str.upper() == 'NO APROBADO'])

                resumen_data.append({
                    'Aprendiz': nombre_completo,
                    'Por Evaluar': por_evaluar,
                    'Aprobadas': aprobado,
                    'No Aprobadas': no_aprobado,
                    'Total': len(datos_aprendiz)
                })

            df_resumen = pd.DataFrame(resumen_data)

            # Mostrar tabla resumen
            st.dataframe(
                df_resumen,
                use_container_width=True,
                height=400,
                hide_index=True
            )

            # Generar documento descargable
            resumen_doc = Document()
            resumen_doc.add_heading("Resumen de Competencias por Aprendiz", level=1)
            resumen_doc.add_paragraph(f"Ficha: {ficha} - Denominación: {denominacion}")
            resumen_doc.add_paragraph(f"Generado el: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
            resumen_doc.add_paragraph(" ")

            # Crear tabla en el documento
            table = resumen_doc.add_table(rows=1, cols=5)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Aprendiz"
            hdr_cells[1].text = "Por Evaluar"
            hdr_cells[2].text = "Aprobadas"
            hdr_cells[3].text = "No Aprobadas"
            hdr_cells[4].text = "Total"

            for _, row in df_resumen.iterrows():
                row_cells = table.add_row().cells
                row_cells[0].text = str(row['Aprendiz'])
                row_cells[1].text = str(row['Por Evaluar'])
                row_cells[2].text = str(row['Aprobadas'])
                row_cells[3].text = str(row['No Aprobadas'])
                row_cells[4].text = str(row['Total'])

            # Botón de descarga
            resumen_bytes = BytesIO()
            resumen_doc.save(resumen_bytes)

            st.download_button(
                label="📥 Descargar Resumen Completo (Word)",
                data=resumen_bytes.getvalue(),
                file_name=f"resumen_aprendices_{ficha}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )


        with pestañas[2]:
            st.subheader("📤 Exportar Reportes Personalizados")

            col_filtro, col_vista = st.columns([1, 2])

            with col_filtro:
                st.markdown("### 🔍 Tipo de Reporte")

                tipo_reporte = st.radio(
                    "Seleccione el tipo de reporte:",
                    ["👥 Reporte por Aprendiz"],
                    help="Elija el tipo de reporte que desea generar"
                )

                if tipo_reporte == "👥 Reporte por Aprendiz":
                    aprendices = df["Nombre"] + " " + df["Apellidos"]
                    aprendiz_seleccionado = st.selectbox(
                        "Seleccionar Aprendiz:",
                        sorted(aprendices.unique()),
                        help="Seleccione un aprendiz para generar su reporte individual"
                    )
                    df_filtrado = df[aprendices == aprendiz_seleccionado]
                else:  # Reporte completo por ficha
                    df_filtrado = df

            with col_vista:
                st.markdown("### 📊 Vista Previa")

                # Mostrar estadísticas básicas
                total_registros = len(df_filtrado)
                total_aprendices = df_filtrado[['Nombre', 'Apellidos']].drop_duplicates().shape[0]

                col1, col2 = st.columns(2)
                with col1:
                    st.metric("Total Registros", total_registros)
                with col2:
                    st.metric("Total Aprendices", total_aprendices)

                # Vista previa de datos
                st.dataframe(
                    df_filtrado.head(8),
                    use_container_width=True,
                    height=280
                )
                if len(df_filtrado) > 8:
                    st.caption(f"Mostrando 8 de {len(df_filtrado)} registros")

            # Generar reporte Word
            doc = Document()

            # Título principal
            doc.add_heading("SENA - Reporte de Juicios Evaluativos", level=1)

            # Información de la ficha
            doc.add_paragraph(f"Ficha: {ficha}")
            doc.add_paragraph(f"Programa: {denominacion}")
            doc.add_paragraph(f"Estado: {estado_ficha}")

            # Información específica del reporte
            if tipo_reporte == "👥 Reporte por Aprendiz":
                doc.add_paragraph(f"Aprendiz: {aprendiz_seleccionado}")

            doc.add_paragraph(f"Fecha de reporte: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
            doc.add_paragraph(" ")

            # Datos detallados - Siempre agrupados por aprendiz
            if tipo_reporte == "👥 Reporte por Aprendiz":
                # Reporte individual del aprendiz
                doc.add_heading(f"Competencias de {aprendiz_seleccionado}", level=2)

                tabla_aprendiz = doc.add_table(rows=len(df_filtrado)+1, cols=3)
                tabla_aprendiz.style = "Table Grid"

                # Encabezados
                tabla_aprendiz.cell(0, 0).text = "Competencia"
                tabla_aprendiz.cell(0, 1).text = "Resultado de Aprendizaje"
                tabla_aprendiz.cell(0, 2).text = "Juicio Evaluativo"

                # Datos
                for i, (_, fila) in enumerate(df_filtrado.iterrows(), 1):
                    tabla_aprendiz.cell(i, 0).text = str(fila['Competencia'])
                    tabla_aprendiz.cell(i, 1).text = str(fila['Resultado de Aprendizaje'])
                    tabla_aprendiz.cell(i, 2).text = str(fila['Juicio de Evaluación'])

            else:
                # Reporte completo de ficha agrupado por aprendiz
                aprendices_agrupados = df_filtrado.groupby(["Nombre", "Apellidos"])

                for (nombre, apellidos), grupo in aprendices_agrupados:
                    # Título del aprendiz
                    doc.add_heading(f"Aprendiz: {nombre} {apellidos}", level=2)

                    # Tabla de competencias del aprendiz
                    tabla_aprendiz = doc.add_table(rows=len(grupo)+1, cols=3)
                    tabla_aprendiz.style = "Table Grid"

                    # Encabezados
                    tabla_aprendiz.cell(0, 0).text = "Competencia"
                    tabla_aprendiz.cell(0, 1).text = "Resultado de Aprendizaje"
                    tabla_aprendiz.cell(0, 2).text = "Juicio Evaluativo"

                    # Datos
                    for i, (_, fila) in enumerate(grupo.iterrows(), 1):
                        tabla_aprendiz.cell(i, 0).text = str(fila['Competencia'])
                        tabla_aprendiz.cell(i, 1).text = str(fila['Resultado de Aprendizaje'])
                        tabla_aprendiz.cell(i, 2).text = str(fila['Juicio de Evaluación'])

                    doc.add_paragraph(" ")

            # Generar nombre del archivo
            if tipo_reporte == "👥 Reporte por Aprendiz":
                nombre_archivo = f"reporte_{aprendiz_seleccionado.replace(' ', '_')}_{ficha}.docx"
            else:
                nombre_archivo = f"reporte_completo_ficha_{ficha}.docx"

            doc_bytes = BytesIO()
            doc.save(doc_bytes)

            st.markdown("---")
            st.markdown("### 📥 Descargar Reporte")

            st.download_button(
                "📄 Descargar Reporte Word",
                data=doc_bytes.getvalue(),
                file_name=nombre_archivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )


        # --- Ayuda ---
        with pestañas[3]:
            st.subheader("❓ Centro de Ayuda")

            col_info, col_soporte = st.columns(2)

            with col_info:
                st.markdown("### 📋 Requisitos del Archivo")
                st.info("""
                **Formato aceptado:** Excel (.xls, .xlsx)

                **Columnas requeridas:**
                - Nombre
                - Apellidos
                - Competencia
                - Resultado de Aprendizaje
                - Juicio de Evaluación

                **Tamaño máximo:** 200MB
                """)

            with col_soporte:
                st.markdown("### 🛠️ Soporte Técnico")
                st.warning("""
                **Problemas comunes:**
                - Formato de archivo incorrecto
                - Columnas faltantes
                - Datos corruptos

                **Contacto:**
                📧 soporte@sena.edu.co
                📞 01-8000-123456
                """)

            st.markdown("---")
            st.markdown("### 🎯 Guía Rápida")
            st.write("1. **Sube tu archivo Excel** con el formato correcto")
            st.write("2. **Revisa las estadísticas** generales en la pestaña principal")
            st.write("3. **Filtra y analiza** datos específicos por aprendiz o competencia")
            st.write("4. **Exporta reportes** en formato Word para tu documentación")

    except Exception as e:
        st.error(f"❌ Error al procesar el archivo: {str(e)}")
        st.info("💡 Verifica que el archivo tenga el formato correcto y todas las columnas requeridas.")
else:
    st.info("""
    **⬆️ Para comenzar, sube tu archivo Excel con los datos de juicios evaluativos.**

    Asegúrate de que el archivo contenga las siguientes columnas:
    - Nombre
    - Apellidos
    - Competencia
    - Resultado de Aprendizaje
    - Juicio de Evaluación
    """)
